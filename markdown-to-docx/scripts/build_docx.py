#!/usr/bin/env python3
from __future__ import annotations

import argparse
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from markdown import markdown

TITLE_SIZE = 20
H1_SIZE = 16
H2_SIZE = 14
H3_SIZE = 12
CAPTION_SIZE = 9
LINK_COLOR = RGBColor(0x05, 0x63, 0xC1)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Convert Markdown to a Word .docx with embedded local images.")
    parser.add_argument("input", help="Markdown source file.")
    parser.add_argument("-o", "--output", help="Output .docx path. Defaults to input basename with .docx suffix.")
    parser.add_argument("--base-dir", help="Base directory for relative image paths. Defaults to input file directory.")
    parser.add_argument("--keep-duplicate-images", action="store_true", help="Keep repeated images.")
    parser.add_argument("--body-font", default="等线", help="Body font. Default: 等线.")
    parser.add_argument("--heading-font", default="黑体", help="Heading font. Default: 黑体.")
    parser.add_argument("--body-size", type=float, default=11, help="Body font size. Default: 11.")
    parser.add_argument(
        "--default-image-width",
        type=float,
        default=5.8,
        help="Default image width in inches when no width attribute is present. Default: 5.8.",
    )
    return parser.parse_args()


def set_run_font(
    run,
    font_name: str,
    size: float,
    *,
    bold: bool | None = None,
    italic: bool | None = None,
    color=None,
    underline: bool | None = None,
) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if underline is not None:
        run.underline = underline


def set_style_font(style, font_name: str, size: float, *, bold: bool = False, italic: bool = False) -> None:
    style.font.name = font_name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def normalize_text(value: str) -> str:
    return value.replace("\xa0", " ")


class Builder:
    def __init__(
        self,
        base_dir: Path,
        *,
        keep_duplicate_images: bool = False,
        body_font: str = "等线",
        heading_font: str = "黑体",
        body_size: float = 11,
        default_image_width: float = 5.8,
    ):
        self.base_dir = base_dir
        self.keep_duplicate_images = keep_duplicate_images
        self.body_font = body_font
        self.heading_font = heading_font
        self.body_size = body_size
        self.default_image_width = default_image_width
        self.seen_images: set[str] = set()
        self.doc = Document()
        section = self.doc.sections[0]
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        self.configure_styles()

    def configure_styles(self) -> None:
        set_style_font(self.doc.styles["Normal"], self.body_font, self.body_size)
        for style_name in ("List Bullet", "List Number", "Quote"):
            set_style_font(self.doc.styles[style_name], self.body_font, self.body_size)
        set_style_font(self.doc.styles["Title"], self.heading_font, TITLE_SIZE, bold=True)
        set_style_font(self.doc.styles["Heading 1"], self.heading_font, H1_SIZE, bold=True)
        set_style_font(self.doc.styles["Heading 2"], self.heading_font, H2_SIZE, bold=True)
        set_style_font(self.doc.styles["Heading 3"], self.heading_font, H3_SIZE, bold=True)

    def add_run(self, paragraph, text: str, *, bold=False, italic=False, link=False) -> None:
        if not text:
            return
        run = paragraph.add_run(text)
        set_run_font(
            run,
            self.body_font,
            self.body_size,
            bold=bold,
            italic=italic,
            color=LINK_COLOR if link else None,
            underline=True if link else None,
        )

    def append_inline(self, paragraph, node, *, bold=False, italic=False, link=False) -> None:
        if isinstance(node, NavigableString):
            self.add_run(paragraph, normalize_text(str(node)), bold=bold, italic=italic, link=link)
            return
        if not isinstance(node, Tag):
            return
        name = node.name.lower()
        if name == "br":
            paragraph.add_run().add_break()
            return
        if name in ("strong", "b"):
            for child in node.children:
                self.append_inline(paragraph, child, bold=True, italic=italic, link=link)
            return
        if name in ("em", "i"):
            for child in node.children:
                self.append_inline(paragraph, child, bold=bold, italic=True, link=link)
            return
        if name == "code":
            for child in node.children:
                self.append_inline(paragraph, child, bold=bold, italic=italic, link=link)
            return
        if name == "a":
            for child in node.children:
                self.append_inline(paragraph, child, bold=bold, italic=italic, link=True)
            return
        if name == "img":
            self.add_image(node)
            return
        for child in node.children:
            self.append_inline(paragraph, child, bold=bold, italic=italic, link=link)

    def width_from_tag(self, img_tag: Tag) -> Inches:
        width_attr = img_tag.get("width")
        if width_attr:
            try:
                width_px = float(str(width_attr).rstrip("px"))
                return Inches(max(2.1, min(6.0, width_px / 96.0)))
            except ValueError:
                pass
        return Inches(max(2.1, min(6.0, self.default_image_width)))

    def add_image(self, img_tag: Tag, caption_override: str | None = None) -> bool:
        src = img_tag.get("src")
        if not src:
            return False
        if src.startswith(("http://", "https://", "data:")):
            paragraph = self.doc.add_paragraph(f"[Unsupported image source: {src}]")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if paragraph.runs:
                set_run_font(paragraph.runs[0], self.body_font, self.body_size, italic=True)
            return False
        if not self.keep_duplicate_images and src in self.seen_images:
            print(f"[skip duplicate image] {src}")
            return False
        img_path = (self.base_dir / src).resolve()
        if not img_path.exists():
            paragraph = self.doc.add_paragraph(f"[Missing image: {src}]")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if paragraph.runs:
                set_run_font(paragraph.runs[0], self.body_font, self.body_size, italic=True)
            return False
        self.seen_images.add(src)
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(img_path), width=self.width_from_tag(img_tag))
        caption = caption_override or img_tag.get("alt", "").strip()
        if caption:
            cp = self.doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = cp.add_run(caption)
            set_run_font(caption_run, self.body_font, CAPTION_SIZE, italic=True)
        return True

    def add_paragraph_from_tag(self, tag: Tag, style=None):
        paragraph = self.doc.add_paragraph(style=style)
        for child in tag.children:
            self.append_inline(paragraph, child)
        return paragraph

    def add_heading(self, text: str, level: int) -> None:
        paragraph = self.doc.add_paragraph(style=f"Heading {level}")
        run = paragraph.add_run(text)
        size = {1: H1_SIZE, 2: H2_SIZE, 3: H3_SIZE}.get(level, H3_SIZE)
        set_run_font(run, self.heading_font, size, bold=True)

    def handle_list(self, tag: Tag, ordered: bool) -> None:
        style = "List Number" if ordered else "List Bullet"
        for li in tag.find_all("li", recursive=False):
            paragraph = self.doc.add_paragraph(style=style)
            for child in li.children:
                if isinstance(child, Tag) and child.name.lower() in ("ul", "ol"):
                    continue
                self.append_inline(paragraph, child)
            for nested in li.find_all(["ul", "ol"], recursive=False):
                self.handle_list(nested, nested.name.lower() == "ol")

    def handle_table(self, table: Tag) -> None:
        for row in table.find_all("tr", recursive=False):
            cells = row.find_all(["td", "th"], recursive=False)
            if not cells:
                continue
            first = cells[0]
            heading_tag = first.find("strong")
            heading_text = " ".join(heading_tag.stripped_strings) if heading_tag else None
            if heading_text:
                self.add_heading(heading_text, level=3)
            desc_parts: list[str] = []
            for child in first.children:
                if isinstance(child, Tag) and child.name.lower() == "strong":
                    continue
                text_part = " ".join(child.stripped_strings) if isinstance(child, Tag) else normalize_text(str(child)).strip()
                if text_part:
                    desc_parts.append(text_part)
            desc = " ".join(desc_parts).strip()
            if desc:
                paragraph = self.doc.add_paragraph()
                run = paragraph.add_run(desc)
                set_run_font(run, self.body_font, self.body_size)
            for cell in cells[1:]:
                img = cell.find("img")
                caption_tag = cell.find("strong")
                caption = " ".join(caption_tag.stripped_strings) if caption_tag else None
                if img:
                    self.add_image(img, caption_override=caption)
                else:
                    raw = " ".join(cell.stripped_strings)
                    if raw:
                        paragraph = self.doc.add_paragraph()
                        run = paragraph.add_run(raw)
                        set_run_font(run, self.body_font, self.body_size)
            self.doc.add_paragraph()

    def build(self, markdown_text: str) -> Document:
        html = markdown(markdown_text, extensions=["tables", "fenced_code", "nl2br", "sane_lists"])
        soup = BeautifulSoup(f"<div>{html}</div>", "lxml")
        root = soup.find("div")
        for child in root.children:
            if isinstance(child, NavigableString):
                if not normalize_text(str(child)).strip():
                    continue
                paragraph = self.doc.add_paragraph()
                run = paragraph.add_run(normalize_text(str(child)).strip())
                set_run_font(run, self.body_font, self.body_size)
                continue
            if not isinstance(child, Tag):
                continue
            name = child.name.lower()
            if name == "h1":
                paragraph = self.doc.add_paragraph(style="Title")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(" ".join(child.stripped_strings))
                set_run_font(run, self.heading_font, TITLE_SIZE, bold=True)
            elif name == "h2":
                self.add_heading(" ".join(child.stripped_strings), level=1)
            elif name == "h3":
                self.add_heading(" ".join(child.stripped_strings), level=2)
            elif name == "h4":
                self.add_heading(" ".join(child.stripped_strings), level=3)
            elif name == "p":
                if child.find("img") and not child.get_text(strip=True):
                    self.add_image(child.find("img"))
                else:
                    self.add_paragraph_from_tag(child)
            elif name == "blockquote":
                self.add_paragraph_from_tag(child, style="Quote")
            elif name == "ul":
                self.handle_list(child, ordered=False)
            elif name == "ol":
                self.handle_list(child, ordered=True)
            elif name == "hr":
                self.doc.add_paragraph()
            elif name == "table":
                self.handle_table(child)
            else:
                self.add_paragraph_from_tag(child)
        for paragraph in self.doc.paragraphs:
            if paragraph.text.strip():
                paragraph.paragraph_format.space_after = Pt(8)
        return self.doc


def main() -> None:
    args = parse_args()
    input_path = Path(args.input).resolve()
    output_path = Path(args.output).resolve() if args.output else input_path.with_suffix(".docx")
    base_dir = Path(args.base_dir).resolve() if args.base_dir else input_path.parent
    markdown_text = input_path.read_text(encoding="utf-8")
    builder = Builder(
        base_dir,
        keep_duplicate_images=args.keep_duplicate_images,
        body_font=args.body_font,
        heading_font=args.heading_font,
        body_size=args.body_size,
        default_image_width=args.default_image_width,
    )
    doc = builder.build(markdown_text)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(output_path)


if __name__ == "__main__":
    main()
