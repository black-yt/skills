#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import Iterable

from docx import Document

ZH_SCALE = 1.14
WORD_RE = re.compile(r"[A-Za-z0-9_][A-Za-z0-9_+\-./<>]*")


def iter_docx_text(doc: Document) -> Iterable[str]:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            yield text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        yield text


def count_text(text: str) -> dict[str, int | str]:
    cjk = sum(1 for ch in text if "\u4e00" <= ch <= "\u9fff")
    words = len(WORD_RE.findall(text))
    nonspace = sum(1 for ch in text if not ch.isspace())
    if cjk > 0:
        calibrated = round((cjk + words) * ZH_SCALE)
        mode = "zh-mixed"
    else:
        calibrated = words
        mode = "en"
    return {
        "mode": mode,
        "cjk": cjk,
        "words": words,
        "nonspace": nonspace,
        "calibrated": calibrated,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Estimate Word-visible counts from generated .docx files.")
    parser.add_argument("paths", nargs="+", help="One or more .docx files.")
    args = parser.parse_args()

    for raw in args.paths:
        path = Path(raw)
        doc = Document(path)
        text = "\n".join(iter_docx_text(doc))
        stats = count_text(text)
        print(path.name)
        print(f"  mode:        {stats['mode']}")
        print(f"  cjk_chars:   {stats['cjk']}")
        print(f"  word_tokens: {stats['words']}")
        print(f"  nonspace:    {stats['nonspace']}")
        print(f"  calibrated:  {stats['calibrated']}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
